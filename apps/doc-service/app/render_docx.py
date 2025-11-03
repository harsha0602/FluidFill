import base64
import io
import re
from typing import Dict, List, Sequence, Tuple

import docx

PatternList = List[Tuple[re.Pattern[str], str]]


UNDERSCORE_RUN_RE = re.compile(r"_{3,}")


def _expand_placeholder_variants(raw_key: str) -> Sequence[str]:
    """Return literal tokens we should replace for a given placeholder key."""
    key = (raw_key or "").strip()
    if not key:
        return []

    variants: List[str] = []

    if key.startswith("[") and key.endswith("]"):
        variants.append(key)
        inner = key[1:-1].strip()
        if inner and UNDERSCORE_RUN_RE.fullmatch(inner):
            variants.append(inner)
        return list(dict.fromkeys(variants))

    if UNDERSCORE_RUN_RE.fullmatch(key):
        return [key]

    if " " in key or "_" in key:
        display = key.replace("_", " ").strip()
        if display:
            variants.append(f"[{display}]")
        return list(dict.fromkeys(variants))

    # Skip bare single-word tokens (e.g., "company") to avoid false positives.
    return []


def _prepare_patterns(mapping: Dict[str, str]) -> PatternList:
    patterns: PatternList = []
    seen: set[str] = set()
    for key, value in mapping.items():
        if value is None:
            continue
        value_str = str(value)
        tokens = list(_expand_placeholder_variants(str(key)))
        tokens.sort(key=lambda tok: (0 if tok.startswith("[") and tok.endswith("]") else 1, -len(tok)))
        for token in tokens:
            token_key = token.strip()
            if not token_key or token_key in seen:
                continue
            seen.add(token_key)
            patterns.append((re.compile(re.escape(token_key), re.IGNORECASE), value_str))
    return patterns


def _replace_in_text(text: str, patterns: PatternList) -> Tuple[str, int]:
    total = 0
    updated = text
    for pattern, replacement in patterns:
        updated, count = pattern.subn(replacement, updated)
        total += count
    return updated, total


def _replace_in_paragraph(paragraph, patterns: PatternList) -> int:
    if not getattr(paragraph, "runs", None):
        return 0
    original = "".join(run.text for run in paragraph.runs)
    if not original:
        return 0
    replaced, count = _replace_in_text(original, patterns)
    if count and replaced != original:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = replaced
    return count


def _replace_in_table(table, patterns: PatternList) -> int:
    total = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                total += _replace_in_paragraph(paragraph, patterns)
            for nested in cell.tables:
                total += _replace_in_table(nested, patterns)
    return total


def render_docx_from_b64(
    doc_b64: str, mapping: Dict[str, str], suggested_name: str | None = None
) -> Tuple[str, int, str]:
    raw = base64.b64decode(doc_b64)
    if len(raw) > 5 * 1024 * 1024:
        raise ValueError("file_too_large")

    patterns = _prepare_patterns(mapping)

    buffer = io.BytesIO(raw)
    document = docx.Document(buffer)

    total_replacements = 0
    for paragraph in document.paragraphs:
        total_replacements += _replace_in_paragraph(paragraph, patterns)
    for table in document.tables:
        total_replacements += _replace_in_table(table, patterns)
    for section in document.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                total_replacements += _replace_in_paragraph(paragraph, patterns)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                total_replacements += _replace_in_paragraph(paragraph, patterns)

    output_stream = io.BytesIO()
    document.save(output_stream)
    encoded = base64.b64encode(output_stream.getvalue()).decode("utf-8")
    base_name = (suggested_name or "document.docx").rsplit(".", 1)[0]
    filled_name = f"{base_name}_filled.docx"
    return encoded, total_replacements, filled_name
